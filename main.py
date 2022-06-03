from genericpath import isfile
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm
from docx import Document
import os

# каталоги проектов, лежащие на одном уровне с текущим проектом
parse_projects = [
    "magneex-backend", "magneex-frontend"
]

# Настройки стилей для приложений: Б, В, Г и т.д бл***
attachment_name = "Б"  
attachment_name_size = 14
attachment_name_first_line_indent_mm = 12.5
attachment_name_line_spacing = 1.5
attachment_name_fontface = 'Times New Roman'

# Настройки стилей для листингов
attachment_code_size = 8
attachment_code_line_spacing = 1
attachment_code_left_line_indent_mm = attachment_name_first_line_indent_mm
attachment_code_fontface = 'Cascadia Mono'

parse_ignore_dirs = [
    "config", "public", "bin", ".git", "vendor", 
    "node_modules", "win32", "Windows", "root", "assets",
    "migrations"
]

parse_ignore_files = [
    "README.md", "package.json", "package-lock.json", ".gitignore",
    "symfony.lock", "symfony.json", "composer.json",
    "composer.lock", "centrifugo-config.json", ".env",
    "php-local.ini"
]

# если комментирование кода и сам код находится хоть где-нибудь в проекте на одной строке
# то лучше воздержаться от использования этой фичи, т.е убрать от сюда все исключения и / или
# оставить только необходимое
parse_ignore_line_content_symbols = [
    "/**", "*", "//", "*/"
]

parser_path = os.getcwd() + "/../"

document = Document()
styles = document.styles

attachemt_main_style = styles.add_style('attachemt_main_title', WD_STYLE_TYPE.PARAGRAPH)
attachemt_main_font = attachemt_main_style.font
attachemt_main_font.bold = True
attachemt_main_font.name = attachment_name_fontface
attachemt_main_font.size = Pt(14)

code_name_style = styles.add_style('code_name', WD_STYLE_TYPE.PARAGRAPH)
code_name_font = code_name_style.font
code_name_font.name = attachment_name_fontface
code_name_font.size = Pt(attachment_name_size)
code_name_format = code_name_style.paragraph_format
code_name_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
code_name_format.first_line_indent = Mm(attachment_name_first_line_indent_mm)
code_name_format.line_spacing = attachment_name_line_spacing

code_body_style = styles.add_style('code_body', WD_STYLE_TYPE.PARAGRAPH)
code_body_font = code_body_style.font
code_body_font.name = attachment_code_fontface
code_body_font.size = Pt(attachment_code_size)
code_body_format = code_body_style.paragraph_format
code_body_format.line_spacing = attachment_code_line_spacing
code_body_format.left_indent = Mm(attachment_code_left_line_indent_mm)

_files_counter = 0

def read_project_files(proj_path, parent_dir = ''):
    proj_files = os.listdir(proj_path)

    for obj in proj_files:
        proj_obj = proj_path + "/" + obj
        
        if os.path.isdir(proj_obj) and obj not in parse_ignore_dirs:
            read_project_files(proj_obj, parent_dir = obj)
        elif os.path.isfile(proj_obj) and obj not in parse_ignore_files:
            file_name = obj

            # помещаем родительский каталог в название файлов, во избежание дубликатов
            if len(parent_dir) > 0:
                file_name = parent_dir + '/' + file_name

            put_to_docx(proj_obj, file_name)
            print("Прочитан файл " + proj_obj)

def put_to_docx(file_path, file_name):
    global _files_counter
    _files_counter = _files_counter + 1
    content_name_paragraph = document.add_paragraph('Листинг ' + attachment_name + '.' + str(_files_counter) + ' – ' + "Файл " + file_name)
    content_name_paragraph.style = styles['code_name']
    file_str_content = ''
    
    try:
        with open(file_path, mode = 'r', encoding="utf8", errors='ignore') as f:
            file_lines = f.readlines()

            for line in file_lines:
                line_is_good = True
                
                for except_content in parse_ignore_line_content_symbols:
                    if except_content in line:
                        line_is_good = False
                
                if line_is_good:
                    file_str_content = file_str_content + line

            content_paragraph = document.add_paragraph(file_str_content)
            content_paragraph.style = styles['code_body']
            f.close()
    except:
        print("Возникла ошибка при чтении файла " + file_path)

def run():
    global _project_counter
    attachment_title = document.add_paragraph('Приложение ' + attachment_name)
    attachment_title.style = styles['attachemt_main_title']
    attachment_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    attachment_subtitle = document.add_paragraph('Документированный листинг программных модулей')
    attachment_subtitle.style = styles['code_name']

    for project in parse_projects:
        proj_path = parser_path + project

        if os.path.exists(proj_path):
            print("Парсинг проекта " + project + " запущен")

            read_project_files(proj_path, parent_dir=project)
        else:
            print("Проекта " + project + " (" + proj_path + ") не существует")
    
    document.add_page_break()
    document.save('attachment_' + attachment_name + '.docx')

run()
